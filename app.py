import streamlit as st
import streamlit_authenticator as stauth
import openai
import os
import sqlite3
import datetime
import requests
import docx
from PyPDF2 import PdfReader
from io import BytesIO
from docx import Document
from fpdf import FPDF
from dotenv import load_dotenv

load_dotenv()

# --- CONFIG ---
openai.api_key = os.getenv("OPENAI_API_KEY")
SERPAPI_KEY = os.getenv("SERPAPI_API_KEY")
DB_PATH = "user_logs.db"

# --- AUTHENTICATION ---
names = ['Admin User', 'Basic User']
usernames = ['admin', 'user']
passwords = ['admin_pass', 'user_pass']
hashed_passwords = stauth.Hasher(passwords).generate()
authenticator = stauth.Authenticate(names, usernames, hashed_passwords, 'cookie', 'secret', cookie_expiry_days=1)

name, auth_status, username = authenticator.login('Login', 'main')
is_admin = username == 'admin'


# Track usage limits in session
if 'usage_count' not in st.session_state:
    st.session_state.usage_count = 0

# Example: limit to 5 API calls per session
if st.session_state.usage_count >= 5:
    st.warning("🚫 Free usage limit reached for this session (5 requests).")
    st.stop()


# --- DB SETUP ---
def init_db():
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("""
        CREATE TABLE IF NOT EXISTS logs (
            username TEXT,
            tool TEXT,
            input TEXT,
            output TEXT,
            timestamp TEXT
        )
    """)
    conn.commit()
    conn.close()

def log_interaction(username, tool, prompt, response):
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("INSERT INTO logs VALUES (?, ?, ?, ?, ?)", 
              (username, tool, prompt, response, datetime.datetime.now().isoformat()))
    conn.commit()
    conn.close()

def fetch_logs(username=None):
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    if username:
        c.execute("SELECT * FROM logs WHERE username = ?", (username,))
    else:
        c.execute("SELECT * FROM logs")
    return c.fetchall()

init_db()

# --- TEMPLATES ---
PROMPT_TEMPLATES = {
    "Academic Summary": "Summarize the following academic article in simple terms:\n\n",
    "SEO Blog": "Write a blog post with SEO optimization about:\n\n",
    "Email Draft": "Draft a professional email to:\n\n"
}

# --- UTILITIES ---
def extract_text(file):
    if file.type == "application/pdf":
        reader = PdfReader(file)
        return "\n".join(page.extract_text() for page in reader.pages)
    elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = docx.Document(file)
        return "\n".join(p.text for p in doc.paragraphs)
    elif file.type == "text/plain":
        return file.read().decode("utf-8")
    return ""

def generate_ai_response(prompt, temperature=0.7):
    response = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[{"role": "user", "content": prompt}],
        temperature=temperature
    )
    return response.choices[0].message["content"]

def web_search(query):
    response = requests.get("https://serpapi.com/search", params={"q": query, "api_key": SERPAPI_KEY})
    data = response.json()
    return "\n\n".join(r['snippet'] for r in data.get("organic_results", []) if 'snippet' in r)

def export_file(content, format_):
    if format_ == "txt":
        return BytesIO(content.encode()), "text/plain"
    elif format_ == "docx":
        buffer = BytesIO()
        doc = Document()
        doc.add_paragraph(content)
        doc.save(buffer)
        return buffer, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif format_ == "pdf":
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        for line in content.split('\n'):
            pdf.cell(200, 10, txt=line, ln=True)
        buffer = BytesIO()
        pdf.output(buffer)
        return buffer, "application/pdf"
    return None, None

# --- MAIN APP ---
if auth_status:
    authenticator.logout("Logout", "sidebar")
    st.sidebar.success(f"Welcome {name}!")

    tab = st.sidebar.radio("Choose a Tool", ["✍️ Write", "🧾 Summarize", "🔍 Research", "🛠️ Edit", "📤 Export", "📜 History"])

    st.title("🧠 AI Writing & Research Assistant")

    uploaded_file = st.file_uploader("Upload a file (PDF, DOCX, TXT)", type=["pdf", "docx", "txt"])
    file_text = extract_text(uploaded_file) if uploaded_file else ""

    prompt_template = st.selectbox("Prompt Template (Optional)", ["None"] + list(PROMPT_TEMPLATES.keys()))
    base_prompt = PROMPT_TEMPLATES.get(prompt_template, "")
    
    prompt_input = st.text_area("Enter your prompt or modify extracted content:", value=base_prompt + file_text, height=300)

    output = ""

    if tab == "✍️ Write":
        if st.session_state.usage_count < 5 and st.button("Generate Content"):
            st.session_state.usage_count += 1
            output = generate_ai_response(prompt_input)
            log_interaction(username, "Write", prompt_input, output)
            st.text_area("Output", output, height=300)

    elif tab == "🧾 Summarize":
        if st.session_state.usage_count < 5 and st.button("Summarize"):
            st.session_state.usage_count += 1
            output = generate_ai_response(f"Summarize this:\n{prompt_input}")
            log_interaction(username, "Summarize", prompt_input, output)
            st.text_area("Summary", output, height=300)

    elif tab == "🔍 Research":
        if st.session_state.usage_count < 5 and st.button("Search"):
            st.session_state.usage_count += 1
            summary = web_search(prompt_input)
            output = generate_ai_response(f"Based on this:\n{summary}")
            log_interaction(username, "Research", prompt_input, output)
            st.text_area("Research Summary", output, height=300)

    elif tab == "🛠️ Edit":
        if st.session_state.usage_count < 5 and st.button("Improve Text"):
            st.session_state.usage_count += 1
            output = generate_ai_response(f"Edit and improve this:\n{prompt_input}")
            log_interaction(username, "Edit", prompt_input, output)
            st.text_area("Improved Text", output, height=300)

    elif tab == "📤 Export":
        text_to_export = st.text_area("Text to Export", prompt_input if not output else output)
        format_ = st.selectbox("Select Format", ["txt", "docx", "pdf"])
        if st.button("Download"):
            buffer, mime = export_file(text_to_export, format_)
            st.download_button("Download", buffer.getvalue(), file_name=f"export.{format_}", mime=mime)

    elif tab == "📜 History":
        st.subheader("Your Past Queries")
        logs = fetch_logs(username if not is_admin else None)
        for user, tool, inp, out, time in logs:
            if is_admin or user == username:
                st.markdown(f"**User:** {user} | **Tool:** {tool} | **Time:** {time}")
                st.markdown(f"**Input:** {inp}")
                st.markdown(f"**Output:** {out}")
                st.markdown("---")

elif auth_status is False:
    st.error("Incorrect username or password.")
elif auth_status is None:
    st.warning("Please enter your credentials.")
